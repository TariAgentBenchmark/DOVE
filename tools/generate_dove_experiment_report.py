#!/usr/bin/env python3

import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
BODY_FONT = "Arial Unicode MS"
CJK_FONT = "Arial Unicode MS"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

PREVIOUS_LIGHTWEIGHT = {
    "psnr": 26.0109,
    "ssim": 0.7675,
    "lpips": 0.2818,
    "dists": 0.1553,
    "clipiqa": 0.4988,
    "raft_warp_l1": 0.01052,
    "frame_diff_l1": 0.02433,
    "steady_seconds": 289.55,
    "speedup": 1.505,
}


def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate the final Chinese DOVE VAE optimization report"
    )
    parser.add_argument("--final_summary", type=Path, required=True)
    parser.add_argument("--screen_metrics", type=Path, required=True)
    parser.add_argument("--screen_temporal", type=Path, required=True)
    parser.add_argument("--visuals", type=Path, nargs="+", required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--final_job", default="2975014")
    parser.add_argument("--report_date", default=date.today().isoformat())
    parser.add_argument(
        "--decoder_checkpoint",
        default=(
            "/data/42-julia-hpc-rz-cv/sig95vg/arb_dove_outputs/perceptual_finish/"
            "perceptual-finish-finish-normalh100-v2_20260730/models/"
            "lpips_continue/checkpoint-4500/decoder.pt"
        ),
    )
    parser.add_argument(
        "--encoder_checkpoint",
        default=(
            "/data/42-julia-hpc-rz-cv/sig95vg/arb_dove_outputs/"
            "vae_encoder_distill/encoder_1111_m1_s3000_2828484/"
            "checkpoint-3000/encoder.pt"
        ),
    )
    return parser.parse_args()


def load_json(path):
    with path.open(encoding="utf-8") as handle:
        return json.load(handle)


def require_final_gates(summary):
    baseline = summary["baseline_metrics"]
    candidate = summary["candidate_metrics"]
    baseline_temporal = summary["baseline_temporal"]
    candidate_temporal = summary["candidate_temporal"]
    gpu_info = " ".join(summary.get("gpu_info", []))
    checks = {
        "正式硬件为 NVIDIA L40S": "l40s" in gpu_info.lower(),
        "同卡无 compile 加速不低于 1.5x": summary["steady_speedup"] >= 1.5,
        "PSNR 不低于 25.5": candidate["psnr"] >= 25.5,
        "LPIPS 不高于原始 DOVE": candidate["lpips"] <= baseline["lpips"],
        "DISTS 不高于原始 DOVE": candidate["dists"] <= baseline["dists"],
        "CLIPIQA 不低于原始 DOVE": candidate["clipiqa"] >= baseline["clipiqa"],
        "RAFT warp 恶化不超过 5%": (
            candidate_temporal["raft_warp_l1"]
            <= baseline_temporal["raft_warp_l1"] * 1.05
        ),
        "帧差误差恶化不超过 5%": (
            candidate_temporal["frame_diff_l1"]
            <= baseline_temporal["frame_diff_l1"] * 1.05
        ),
    }
    failed = [name for name, passed in checks.items() if not passed]
    if failed:
        raise RuntimeError("Final acceptance gates failed: " + ", ".join(failed))
    return checks


def cross_check_screen(summary, screen_metrics, screen_temporal):
    candidate = summary["candidate_metrics"]
    candidate_temporal = summary["candidate_temporal"]
    for key in ("psnr", "ssim", "lpips", "dists", "clipiqa"):
        if abs(candidate[key] - screen_metrics["average"][key]) > 0.01:
            raise RuntimeError(f"Final/screen metric mismatch for {key}")
    for key in ("raft_warp_l1", "frame_diff_l1"):
        if abs(candidate_temporal[key] - screen_temporal["average"][key]) > 0.001:
            raise RuntimeError(f"Final/screen temporal metric mismatch for {key}")


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = BODY_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_text(
    paragraph,
    text,
    *,
    size=11,
    bold=False,
    color=INK,
    italic=False,
    align=None,
):
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, header=True, first_col_left=True):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            if header and row_index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if first_col_left and col_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9,
                        bold=(header and row_index == 0),
                        color=INK,
                    )


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths_dxa)
    style_table(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, *, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, size=11, bold=True, color=INK)
        remainder = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(remainder, size=11, color=INK)
    else:
        set_paragraph_text(paragraph, text)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    set_paragraph_text(paragraph, text)
    return paragraph


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(label + "  ")
    set_run_font(label_run, size=11, bold=True, color=DARK_BLUE)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def set_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    for field_name in ("PAGE",):
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = field_name
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        value = OxmlElement("w:t")
        value.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, instruction, separate, value, end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_text(
        header_paragraph,
        "DOVE 单卡 VAE 纯算法优化",
        size=9,
        bold=True,
        color=MUTED,
    )
    footer = section.footer
    set_page_field(footer.paragraphs[0])


def add_cover(doc, report_date, final_job, summary):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(24)
    paragraph.paragraph_format.space_after = Pt(4)
    set_paragraph_text(
        paragraph,
        "实验报告",
        size=11,
        bold=True,
        color=BLUE,
    )
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    set_paragraph_text(
        title,
        "DOVE 单卡 VAE 解码链路优化",
        size=25,
        bold=True,
        color=INK,
    )
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    set_paragraph_text(
        subtitle,
        "轻量 Encoder/Decoder 蒸馏、感知指标补偿与同卡速度验收",
        size=13,
        color=MUTED,
    )
    metadata = [
        ("基座模型", "DOVE Stage-2"),
        ("最终结构", "Encoder [1,1,1,1] + mid=1；Decoder [1,1,1,2]"),
        ("验证数据", "UDM10（10 段视频，32 帧，4 倍超分）"),
        ("最终验收任务", f"Slurm job {final_job}，单 NVIDIA L40S"),
        ("报告日期", report_date),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(label + "：")
        set_run_font(label_run, size=10.5, bold=True, color=INK)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_callout(
        doc,
        "报告结论",
        (
            "正式 L40S 同卡、双方关闭 torch.compile 的稳态端到端耗时由 "
            f"{summary['baseline_steady_seconds']:.2f} 秒降至 "
            f"{summary['candidate_steady_seconds']:.2f} 秒，达到 "
            f"{summary['steady_speedup']:.3f}x；LPIPS、DISTS、CLIPIQA "
            "均不低于原始 DOVE，全部验收项通过。"
        ),
        fill="EAF5EF",
    )
    doc.add_page_break()


def build_report(args, summary, screen_metrics, screen_temporal, checks):
    doc = Document()
    configure_document(doc)
    add_cover(doc, args.report_date, args.final_job, summary)

    baseline = summary["baseline_metrics"]
    candidate = summary["candidate_metrics"]
    baseline_temporal = summary["baseline_temporal"]
    candidate_temporal = summary["candidate_temporal"]

    add_heading(doc, "结论摘要", 1)
    add_callout(
        doc,
        "最终结论",
        (
            f"在单张 L40S、原始 DOVE 与候选均关闭 torch.compile 的同卡口径下，"
            f"稳态端到端耗时由 {summary['baseline_steady_seconds']:.2f} 秒降至 "
            f"{summary['candidate_steady_seconds']:.2f} 秒，加速 "
            f"{summary['steady_speedup']:.3f}x。最终候选的 LPIPS、DISTS、CLIPIQA "
            "均不低于原始 DOVE，且时序误差保持在 5% 限制内。"
        ),
        fill="EAF5EF",
    )
    add_body(
        doc,
        (
            "本轮没有改变 DOVE 的 one-step latent Transformer，也没有使用编译、"
            "TensorRT、多卡推理、减帧或降低输出分辨率。加速来自 VAE encoder 和 "
            "decoder 的结构裁剪；画质补偿来自 encoder-matched trajectory cache "
            "及以 GT 为目标的 LPIPS/DISTS/CLIPIQA 定向微调。"
        ),
    )

    add_heading(doc, "1. 目标与验收口径", 1)
    add_body(
        doc,
        "客户要求是在不计入 torch.compile 的前提下，将单卡推理速度提升到至少 1.5 倍，同时允许 PSNR/SSIM 小幅下降，但感知类指标不得低于原始 DOVE。",
    )
    gate_rows = []
    for name, passed in checks.items():
        gate_rows.append((name, "通过" if passed else "未通过"))
    add_table(doc, ("验收项", "结果"), gate_rows, [7200, 2160])

    add_heading(doc, "2. 方法", 1)
    add_heading(doc, "2.1 结构裁剪", 2)
    add_body(
        doc,
        "Encoder 的四个 down block 均保留 1 层，mid block 保留 1 层；Decoder 的四个 up block 使用 [1,1,1,2]，在高分辨率端多保留一层。Transformer、scheduler、文本编码器及采样步数保持原始 DOVE 配置。",
    )
    add_heading(doc, "2.2 Encoder-matched trajectory cache", 2)
    add_body(
        doc,
        "Decoder 微调没有直接复用原始 encoder 产生的 latent。先用最终轻量 encoder 对 1024 条真实退化 LQ 条件重新编码，再通过完整 DOVE 推理轨迹得到 decoder 输入 latent；监督目标同时保存原始 DOVE decoder 输出和对应 GT。缓存结果为 1024/1024，失败 0。",
    )
    add_heading(doc, "2.3 感知补偿目标", 2)
    loss_rows = [
        ("Teacher L1", "0.15", "保持与原始 DOVE 输出一致"),
        ("Teacher MSE", "0.03", "抑制整体像素漂移"),
        ("GT L1", "0.08", "补偿 encoder 裁剪带来的偏差"),
        ("Frame difference", "0.10", "约束相邻帧变化"),
        ("DISTS", "0.15", "直接优化结构感知距离"),
        ("LPIPS", "0.25", "直接优化深层感知距离"),
        ("CLIPIQA hinge", "0.01", "维持无参考感知质量"),
    ]
    add_table(doc, ("Loss", "权重", "作用"), loss_rows, [2640, 1440, 5280])
    add_body(
        doc,
        "最终分支从 perceptual checkpoint-4400 继续训练 100 step，学习率 1e-5，在 checkpoint-4500 达到全部感知门槛。",
    )

    doc.add_page_break()
    add_heading(doc, "3. 实验过程", 1)
    process_rows = [
        ("基线轻量版", "Decoder 4100 + Encoder 3000", "1.505x，但 LPIPS/DISTS/CLIPIQA 未达客户新要求"),
        ("第一轮感知 sweep", "3 组 × 200 step", "CLIPIQA 已超过 DOVE；LPIPS/DISTS 仍有差距"),
        ("二次 refine", "3 组 × 100 step", "DISTS 与 CLIPIQA 达标；LPIPS 距门槛 0.0005"),
        ("收尾 continuation", "2 组 × 100 step", "checkpoint-4500 的三项感知指标全部达标"),
        ("最终同卡验收", f"job {args.final_job}", "原始 DOVE 与候选在同一张 L40S 顺序执行"),
    ]
    add_table(doc, ("阶段", "配置", "结果"), process_rows, [2160, 2400, 4800])

    add_heading(doc, "4. 实验设置", 1)
    settings_rows = [
        ("基座模型", "DOVE Stage-2"),
        ("训练数据", "HQ-VSR，1024 条 encoder-matched trajectory cache"),
        ("验证数据", "UDM10，10 段视频"),
        ("输入/输出", "32 帧 180x318 -> 32 帧 720x1272"),
        ("精度", "bfloat16"),
        ("速度硬件", "单 NVIDIA L40S"),
        ("排除项", "torch.compile、TensorRT、CUDA Graph、多卡、减帧、降分辨率"),
    ]
    add_table(doc, ("项目", "设置"), settings_rows, [2400, 6960])

    add_heading(doc, "5. 结果", 1)
    add_heading(doc, "5.1 同卡速度", 2)
    speed_rows = [
        (
            "原始 DOVE",
            f"{summary['baseline_steady_seconds']:.2f}",
            "1.000x",
            f"{summary['baseline_core_seconds']:.2f}",
        ),
        (
            "最终轻量 VAE",
            f"{summary['candidate_steady_seconds']:.2f}",
            f"{summary['steady_speedup']:.3f}x",
            f"{summary['candidate_core_seconds']:.2f}",
        ),
    ]
    add_table(
        doc,
        ("模型", "稳态端到端（秒）", "加速", "核心推理（秒）"),
        speed_rows,
        [2880, 2400, 1680, 2400],
    )
    add_body(
        doc,
        "稳态端到端时间按 run_wall - model_load - model_to_device 计算，包含预处理、Transformer、VAE encode/decode、输出拼接及视频保存。",
    )

    add_heading(doc, "5.2 画质", 2)
    quality_rows = [
        (
            "原始 DOVE",
            f"{baseline['psnr']:.4f}",
            f"{baseline['ssim']:.4f}",
            f"{baseline['lpips']:.4f}",
            f"{baseline['dists']:.4f}",
            f"{baseline['clipiqa']:.4f}",
        ),
        (
            "此前轻量版",
            f"{PREVIOUS_LIGHTWEIGHT['psnr']:.4f}",
            f"{PREVIOUS_LIGHTWEIGHT['ssim']:.4f}",
            f"{PREVIOUS_LIGHTWEIGHT['lpips']:.4f}",
            f"{PREVIOUS_LIGHTWEIGHT['dists']:.4f}",
            f"{PREVIOUS_LIGHTWEIGHT['clipiqa']:.4f}",
        ),
        (
            "最终候选",
            f"{candidate['psnr']:.4f}",
            f"{candidate['ssim']:.4f}",
            f"{candidate['lpips']:.4f}",
            f"{candidate['dists']:.4f}",
            f"{candidate['clipiqa']:.4f}",
        ),
    ]
    add_table(
        doc,
        ("模型", "PSNR ↑", "SSIM ↑", "LPIPS ↓", "DISTS ↓", "CLIPIQA ↑"),
        quality_rows,
        [2400, 1392, 1392, 1392, 1392, 1392],
    )
    add_body(
        doc,
        (
            f"相较此前轻量版，最终候选的 LPIPS 改善 "
            f"{(1 - candidate['lpips'] / PREVIOUS_LIGHTWEIGHT['lpips']) * 100:.2f}%，"
            f"DISTS 改善 {(1 - candidate['dists'] / PREVIOUS_LIGHTWEIGHT['dists']) * 100:.2f}%，"
            f"CLIPIQA 提升 {(candidate['clipiqa'] / PREVIOUS_LIGHTWEIGHT['clipiqa'] - 1) * 100:.2f}%。"
            "PSNR/SSIM 仍略低于原始 DOVE，但处于客户允许范围。"
        ),
    )

    add_heading(doc, "5.3 时序", 2)
    temporal_rows = [
        (
            "原始 DOVE",
            f"{baseline_temporal['raft_warp_l1']:.6f}",
            f"{baseline_temporal['frame_diff_l1']:.6f}",
        ),
        (
            "最终候选",
            f"{candidate_temporal['raft_warp_l1']:.6f}",
            f"{candidate_temporal['frame_diff_l1']:.6f}",
        ),
    ]
    add_table(
        doc,
        ("模型", "RAFT warp L1 ↓", "Frame-difference L1 ↓"),
        temporal_rows,
        [3600, 2880, 2880],
    )
    add_body(
        doc,
        "两项时序误差均低于原始 DOVE 的 1.05 倍上限。四段视频、三个时间点的视觉抽查未发现新增闪烁、错位或结构跳变。",
    )

    add_heading(doc, "6. 视觉对比", 1)
    add_body(
        doc,
        "下图按列给出原始 DOVE、最终轻量 VAE 和 GT，按行给出第 5、15、25 帧。重点观察人物/物体轮廓、文字区域、重复纹理和相邻时间点的一致性。",
    )
    for index, visual in enumerate(args.visuals, start=1):
        if not visual.is_file():
            raise FileNotFoundError(visual)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        set_paragraph_text(
            paragraph,
            f"图 {index}  UDM10 样例 {visual.stem.split('_')[0]}",
            size=9.5,
            bold=True,
            color=MUTED,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_after = Pt(8)
        image_paragraph.add_run().add_picture(str(visual), width=Inches(6.45))
        if index < len(args.visuals):
            doc.add_page_break()

    add_heading(doc, "7. 结论与边界", 1)
    add_body(
        doc,
        "本次最终方案可以表述为：在单张 L40S、双方关闭 compile 的同卡测试中，通过轻量 VAE encoder/decoder 蒸馏达到至少 1.5 倍稳态端到端加速；LPIPS、DISTS 和 CLIPIQA 均不低于原始 DOVE，PSNR/SSIM 有小幅下降。",
    )
    add_bullet(doc, "方案不是去 VAE；保留了轻量 encoder 和 decoder。")
    add_bullet(doc, "当前结论来自 UDM10 的 720x1272 输出，不等同于原生 4K 验收。")
    add_bullet(doc, "真实部署仍需复测长视频、4K tile 边界、峰值显存和不同驱动环境。")
    add_bullet(doc, "1.5x 是同卡实验结论，不应直接写成跨机器的固定 SLA。")

    add_heading(doc, "附录 A：最终产物", 1)
    paragraph = add_body(doc, "Decoder checkpoint：")
    paragraph.paragraph_format.keep_with_next = True
    add_body(doc, args.decoder_checkpoint)
    paragraph = add_body(doc, "Encoder checkpoint：")
    paragraph.paragraph_format.keep_with_next = True
    add_body(doc, args.encoder_checkpoint)
    jobs_rows = [
        ("2921911", "1024 条 encoder-matched trajectory cache"),
        ("2921912 / 2921913", "第一轮感知 sweep 与全量评估"),
        ("2974916 / 2974917", "二次 refine 与 6 个 checkpoint 评估"),
        ("2974989 / 2975000", "收尾 continuation 与 4 个 checkpoint 评估"),
        (args.final_job, "单 L40S 同卡最终验收"),
    ]
    add_table(doc, ("Slurm job", "用途"), jobs_rows, [2640, 6720])

    add_heading(doc, "附录 B：筛选结果交叉核对", 1)
    screen_rows = [
        ("PSNR", f"{screen_metrics['average']['psnr']:.4f}"),
        ("SSIM", f"{screen_metrics['average']['ssim']:.4f}"),
        ("LPIPS", f"{screen_metrics['average']['lpips']:.4f}"),
        ("DISTS", f"{screen_metrics['average']['dists']:.4f}"),
        ("CLIPIQA", f"{screen_metrics['average']['clipiqa']:.4f}"),
        ("RAFT warp L1", f"{screen_temporal['average']['raft_warp_l1']:.6f}"),
        ("Frame-difference L1", f"{screen_temporal['average']['frame_diff_l1']:.6f}"),
    ]
    add_table(doc, ("指标", "收尾筛选值"), screen_rows, [4680, 4680])

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)


def main():
    args = parse_args()
    summary = load_json(args.final_summary)
    screen_metrics = load_json(args.screen_metrics)
    screen_temporal = load_json(args.screen_temporal)
    checks = require_final_gates(summary)
    cross_check_screen(summary, screen_metrics, screen_temporal)
    build_report(args, summary, screen_metrics, screen_temporal, checks)
    print(f"Report written to {args.output}")


if __name__ == "__main__":
    main()
